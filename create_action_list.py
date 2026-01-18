from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

# Set narrow margins
for section in doc.sections:
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)

# Title
title = doc.add_heading('CEPHO.Ai Action List', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.color.rgb = RGBColor(0, 0, 0)

# Subtitle
subtitle = doc.add_paragraph('18 January 2026')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Summary box
summary = doc.add_paragraph()
summary.add_run('Total Actions: ').bold = True
summary.add_run('140 items | ')
summary.add_run('Status: ').bold = True
summary.add_run('In Progress | ')
summary.add_run('Generated: ').bold = True
summary.add_run('14:00 GMT+4')

doc.add_paragraph()

# All actions in order
all_actions = [
    # Section 1: UI/UX Review Actions (1-29)
    'Apply consistent design language across all pages (black bg, white text, pink accents)',
    'Replace grey Sepho logo with official pink brain logo on all pages',
    'Standardize terminology: "Digital Twin" → "Chief of Staff", "Experts" → "AI SMEs"',
    'Remove personalized greetings (no "Good morning, John")',
    'Change all grey text to white for readability',
    'Rename "Listen to brief" to "Morning Brief" or "Daily Brief"',
    'Attribute brief to "Victoria" with her picture',
    'Add clear buttons: "Watch the Brief" and "Listen to the Brief"',
    'Evening Review: Apply consistent design (black banner, black background)',
    'Evening Review: Optimize layout to be more compact, reduce scrolling',
    'Evening Review: Add icons for actions (tick, defer, cross)',
    'Chief of Staff: Create clear landing page dashboard showing all action status',
    'Chief of Staff: Make Development Pathway collapsible/hidden by default',
    'Chief of Staff: Fix layout overflow (not extending too far right)',
    'Chief of Staff: Display action status with percentage completion',
    'Chief of Staff: Add Future Development section for ideas tracking',
    'Chief of Staff: Remove/relocate secondary bar (government funding, subscriptions)',
    'Chief of Staff: Add direct chat function',
    'Chief of Staff: Establish as central hub for status updates',
    'AI SMEs: Add visual space between "Recent Chats" and "Browse by type"',
    'AI SMEs: Change SME names text color to white',
    'AI SMEs: Display full SME name when selected (not just icon)',
    'AI SMEs: Fix "Projects" and "Insights" label text colors',
    'AI SMEs: Style chat button in white',
    'Projects/Project Genesis: Apply consistent black/white/pink theme',
    'Innovation Hub: Apply consistent theme, clear layout',
    'Workflow: Fix all formatting to new standards',
    'Knowledge and Vault: Clarify Library vs Document Library distinction',
    'The Vault: Fix text readability (visible colors)',
    
    # Section 2: SME Panel Recommendations (30-49)
    'Dr Sarah Chen: Implement edge caching for global latency reduction',
    'Marcus Webb: Add model versioning and A/B testing framework',
    'Elena Rodriguez: Implement zero-trust architecture for API layer',
    'Thomas Eriksson: Optimize database queries for scale',
    'Lisa Park: Conduct usability testing with 20+ users monthly',
    'David Müller: Reduce cognitive load in complex workflows',
    'Sophie Laurent: Strengthen brand consistency in micro-interactions',
    'Michael Santos: Document component library for consistency',
    'Michael Torres: Prepare investor deck for Series A readiness',
    'Dr Helen Wright: Establish monthly competitor tracking dashboard',
    'Jonathan Blake: Identify potential acquisition targets in AI space',
    'Carlos Mendez: Implement continuous improvement framework',
    'Diana Patel: Create monthly investor update cadence',
    'Emma Wilson: Implement proactive health scoring for accounts',
    'Rachel Green: Launch NPS program with quarterly surveys',
    'Daniel Brown: Build customer community with 1000+ members',
    'William Taylor: Implement churn prediction model',
    'Patrick Sullivan: Implement GDPR compliance automation',
    'Natalie Chen: Conduct annual privacy impact assessments',
    'Andrew Walsh: Establish board advisory committee',
    
    # Section 3: CEPHO Template Repository (50-64)
    'CEO KPI Scorecard Report Template',
    'Chief of Staff Briefing Paper Template',
    'SME Panel Review Report Template',
    'Project Genesis Assessment Template',
    'Morning Signal Daily Briefing Template',
    'Investor Update Template',
    'Board Meeting Pack Template',
    'Strategic Review Template',
    'Competitive Analysis Template',
    'Financial Summary Template',
    'Customer Success Report Template',
    'Development Action List Template',
    'Process Documentation Template',
    'Meeting Minutes Template',
    'Decision Log Template',
    
    # Section 4: Notter Integration (65-70)
    'Research Notter API capabilities',
    'Set up Notter account connection',
    'Build voice note import from Notter',
    'Create sync mechanism for voice notes',
    'Add Notter as fallback when Manus unavailable',
    'Process Notter voice notes through COS pipeline',
    
    # Section 5: Morning Signal System (71-102)
    'Connect ElevenLabs TTS to daily signal content',
    'Create voice briefing script generator',
    'Add voice note player to Morning Signal page',
    'Store generated audio files',
    'Connect AI avatar video generation to daily signal',
    'Create video script from signal content',
    'Generate avatar video presentation',
    'Add video player to Morning Signal page',
    'Generate branded PDF using CEPHO design guidelines',
    'Include all signal sections in PDF format',
    'Add download button to Morning Signal page',
    'Store generated PDFs',
    'All outputs pass through Chief of Staff validation',
    'Check design guidelines compliance',
    'Verify content quality before delivery',
    'Log all outputs for COS learning',
    'Ensure all parts work together seamlessly end-to-end',
    'Single trigger generates all three formats (voice, video, PDF)',
    'Consistent content across all formats',
    'Unified delivery mechanism',
    'Error handling across the pipeline',
    'Progress tracking for multi-format generation',
    'Fallback if one format fails (others still deliver)',
    'Define CEPHO voice tone profile',
    'Create tone guidelines for different content types',
    'Build tone matching system against user preferences',
    'Implement urgency detection and appropriate delivery',
    'Avoid monotone - vary pace and emphasis naturally',
    'Create script structure templates (hook, body, close)',
    'Define attention span guidelines',
    'Build engagement hooks library',
    'Implement boredom prevention',
    
    # Section 6: Content Quality SME Panel (103-114)
    'Script Writer SME - review all scripts before generation',
    'Film Director SME - advise on visual pacing and imagery',
    'Voice Director SME - guide tone and delivery',
    'Engagement Expert SME - optimize for attention retention',
    'Optimize ElevenLabs voice settings for CEPHO tone',
    'Implement dynamic pacing in voice scripts',
    'Add emphasis markers for key points',
    'Create voice style presets (briefing, urgent, casual)',
    'Test and refine voice output quality',
    'Define visual pacing guidelines',
    'Create imagery prompts for key moments',
    'Build transition and visual hook library',
    
    # Section 7: Quality Assurance System (115-124)
    'Define engagement metrics for each format',
    'Build feedback loop for continuous improvement',
    'Track attention retention data',
    'A/B test different approaches',
    'COS reviews all outputs against quality standards',
    'Daily Signal (voice, video, PDF) quality check',
    'Project briefings quality check',
    'SME consultations quality check',
    'Reports and presentations quality check',
    'All external communications quality check',
    
    # Section 8: Digital Twin Questionnaire (125-134)
    'Add 200-question Digital Twin questionnaire to onboarding',
    'Create database table to store questionnaire responses',
    'Build questionnaire UI component with progress tracking',
    'Split questionnaire into digestible sections',
    'Store responses and calculate user profile scores',
    'Generate Digital Twin profile from questionnaire responses',
    'Show COS Understanding Level based on completion',
    'Allow users to update answers over time',
    'Use questionnaire data to personalize COS behavior',
    'Make questionnaire completion part of Getting to 100 score',
    
    # Section 9: Chief of Staff KPI Dashboard (135-140)
    'Create ChiefOfStaffKPIReport component',
    'Build SME Performance Heatmap with color-coded scoring',
    'Create Business Area Scoring Table with all SME rankings',
    'Implement Enhancement Recommendations engine',
    'Add export functionality for report (PDF/CSV)',
    'Integrate report into Chief of Staff page',
]

# Section boundaries
sections = [
    ('1. UI/UX Review Actions', 0, 29),
    ('2. SME Panel Recommendations', 29, 49),
    ('3. CEPHO Template Repository', 49, 64),
    ('4. Notter Integration', 64, 70),
    ('5. Morning Signal System', 70, 102),
    ('6. Content Quality SME Panel', 102, 114),
    ('7. Quality Assurance System', 114, 124),
    ('8. Digital Twin Questionnaire', 124, 134),
    ('9. Chief of Staff KPI Dashboard', 134, 140),
]

for section_name, start, end in sections:
    heading = doc.add_heading(section_name, level=1)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
    
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    
    # Header row
    hdr = table.rows[0].cells
    headers = ['#', 'Action', 'Status', 'Comments']
    for i, h in enumerate(headers):
        hdr[i].text = h
        hdr[i].paragraphs[0].runs[0].bold = True
        hdr[i].paragraphs[0].runs[0].font.size = Pt(9)
    
    # Data rows
    for i in range(start, end):
        row = table.add_row().cells
        row[0].text = str(i + 1)
        row[1].text = all_actions[i]
        row[2].text = '☐'
        row[3].text = ''
        for cell in row:
            cell.paragraphs[0].runs[0].font.size = Pt(9)
    
    doc.add_paragraph()

# Summary Table
doc.add_heading('Summary', level=1)
summary_table = doc.add_table(rows=1, cols=3)
summary_table.style = 'Table Grid'
sum_hdr = summary_table.rows[0].cells
for i, h in enumerate(['Section', 'Items', 'Range']):
    sum_hdr[i].text = h
    sum_hdr[i].paragraphs[0].runs[0].bold = True

sections_summary = [
    ('1. UI/UX Review Actions', '29', '1-29'),
    ('2. SME Panel Recommendations', '20', '30-49'),
    ('3. CEPHO Template Repository', '15', '50-64'),
    ('4. Notter Integration', '6', '65-70'),
    ('5. Morning Signal System', '32', '71-102'),
    ('6. Content Quality SME Panel', '12', '103-114'),
    ('7. Quality Assurance System', '10', '115-124'),
    ('8. Digital Twin Questionnaire', '10', '125-134'),
    ('9. Chief of Staff KPI Dashboard', '6', '135-140'),
    ('TOTAL', '140', '1-140'),
]

for section in sections_summary:
    row = summary_table.add_row().cells
    for i, val in enumerate(section):
        row[i].text = val
    if section[0] == 'TOTAL':
        for cell in row:
            cell.paragraphs[0].runs[0].bold = True

doc.add_paragraph()

# Footer
footer = doc.add_paragraph()
footer.add_run('Document Generated: ').bold = True
footer.add_run('18 January 2026 | 14:00 GMT+4\n')
footer.add_run('Logged by: ').bold = True
footer.add_run('Chief of Staff\n')
footer.add_run('Design Principles: ').bold = True
footer.add_run('Black text (#000000), Pink accents (#E91E8C), NO light grey')

doc.save('/home/ubuntu/the-brain/CEPHO_ACTION_LIST_18JAN2026.docx')
print('Action List document created with sequential numbering 1-140')
