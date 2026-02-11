from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

# Title
title = doc.add_heading('SME Panel Review Actions', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph('To Be Completed Today - 18 January 2026')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

# Summary
summary = doc.add_paragraph()
summary.add_run('Total Actions Selected: ').bold = True
summary.add_run('107 items\n')
summary.add_run('Source: ').bold = True
summary.add_run('CEO KPI Report SME Recommendations + Development Pathway\n')
summary.add_run('Status: ').bold = True
summary.add_run('Logged for Implementation')

doc.add_paragraph()

# Section 1: SME Panel Actions (Original 20)
doc.add_heading('Section 1: SME Panel Actions (20 items)', level=1)

table = doc.add_table(rows=1, cols=6)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
headers = ['#', 'SME Reviewer', 'Panel', 'Role', 'Score', 'Recommendation']
for i, header in enumerate(headers):
    hdr_cells[i].text = header
    hdr_cells[i].paragraphs[0].runs[0].bold = True

actions = [
    ('1', 'Dr Sarah Chen', 'Technology', 'Cloud Architecture', '88', 'Implement edge caching for global latency reduction'),
    ('2', 'Marcus Webb', 'Technology', 'AI Systems', '85', 'Add model versioning and A/B testing framework'),
    ('3', 'Elena Rodriguez', 'Technology', 'Security', '86', 'Implement zero-trust architecture for API layer'),
    ('6', 'Thomas Eriksson', 'Technology', 'Backend', '89', 'Optimize database queries for scale'),
    ('9', 'Lisa Park', 'UI/UX', 'UX Research', '84', 'Conduct usability testing with 20+ users monthly'),
    ('10', 'David Müller', 'UI/UX', 'Interaction Design', '86', 'Reduce cognitive load in complex workflows'),
    ('13', 'Sophie Laurent', 'UI/UX', 'Visual Design', '85', 'Strengthen brand consistency in micro-interactions'),
    ('16', 'Michael Santos', 'UI/UX', 'Design Systems', '86', 'Document component library for consistency'),
    ('20', 'Michael Torres', 'Strategy', 'Investment', '70', 'Prepare investor deck for Series A readiness'),
    ('21', 'Dr Helen Wright', 'Strategy', 'Competitive Intel', '72', 'Establish monthly competitor tracking dashboard'),
    ('22', 'Jonathan Blake', 'Strategy', 'M&A Advisory', '69', 'Identify potential acquisition targets in AI space'),
    ('24', 'Carlos Mendez', 'Operations', 'Process Engineer', '86', 'Implement continuous improvement framework'),
    ('32', 'Diana Patel', 'Finance', 'Investor Relations', '72', 'Create monthly investor update cadence'),
    ('35', 'Emma Wilson', 'Customer', 'Customer Success', '68', 'Implement proactive health scoring for accounts'),
    ('37', 'Rachel Green', 'Customer', 'Customer Research', '70', 'Launch NPS program with quarterly surveys'),
    ('38', 'Daniel Brown', 'Customer', 'Community', '72', 'Build customer community with 1000+ members'),
    ('40', 'William Taylor', 'Customer', 'Retention', '69', 'Implement churn prediction model'),
    ('42', 'Patrick Sullivan', 'Legal', 'Compliance', '65', 'Implement GDPR compliance automation'),
    ('43', 'Natalie Chen', 'Legal', 'Privacy', '70', 'Conduct annual privacy impact assessments'),
    ('44', 'Andrew Walsh', 'Legal', 'Corporate', '68', 'Establish board advisory committee'),
]

for action in actions:
    row_cells = table.add_row().cells
    for i, val in enumerate(action):
        row_cells[i].text = val

doc.add_paragraph()

# Section 2: Morning Signal System
doc.add_heading('Section 2: Morning Signal System (32 items)', level=1)

table2 = doc.add_table(rows=1, cols=4)
table2.style = 'Table Grid'
hdr2 = table2.rows[0].cells
for i, h in enumerate(['#', 'Action', 'Priority', 'Status']):
    hdr2[i].text = h
    hdr2[i].paragraphs[0].runs[0].bold = True

morning_signal = [
    ('60', 'Connect ElevenLabs TTS to daily signal content', 'P1', 'Pending'),
    ('61', 'Create voice briefing script generator', 'P1', 'Pending'),
    ('62', 'Add voice note player to Morning Signal page', 'P1', 'Pending'),
    ('63', 'Store generated audio files', 'P2', 'Pending'),
    ('64', 'Connect AI avatar video generation to daily signal', 'P2', 'Pending'),
    ('65', 'Create video script from signal content', 'P2', 'Pending'),
    ('66', 'Generate avatar video presentation', 'P2', 'Pending'),
    ('67', 'Add video player to Morning Signal page', 'P2', 'Pending'),
    ('68', 'Generate branded PDF using CEPHO design guidelines', 'P1', 'Pending'),
    ('69', 'Include all signal sections in PDF format', 'P1', 'Pending'),
    ('70', 'Add download button to Morning Signal page', 'P1', 'Pending'),
    ('71', 'Store generated PDFs', 'P2', 'Pending'),
    ('72', 'All outputs pass through Chief of Staff validation', 'P1', 'Pending'),
    ('73', 'Check design guidelines compliance', 'P1', 'Pending'),
    ('74', 'Verify content quality before delivery', 'P1', 'Pending'),
    ('75', 'Log all outputs for COS learning', 'P2', 'Pending'),
    ('76', 'Ensure all parts work together seamlessly end-to-end', 'P1', 'Pending'),
    ('77', 'Single trigger generates all three formats (voice, video, PDF)', 'P2', 'Pending'),
    ('78', 'Consistent content across all formats', 'P2', 'Pending'),
    ('79', 'Unified delivery mechanism', 'P2', 'Pending'),
    ('80', 'Error handling across the pipeline', 'P2', 'Pending'),
    ('81', 'Progress tracking for multi-format generation', 'P3', 'Pending'),
    ('82', 'Fallback if one format fails (others still deliver)', 'P2', 'Pending'),
    ('83', 'Define CEPHO voice tone profile', 'P1', 'Pending'),
    ('84', 'Create tone guidelines for different content types', 'P2', 'Pending'),
    ('85', 'Build tone matching system against user preferences', 'P2', 'Pending'),
    ('86', 'Implement urgency detection and appropriate delivery', 'P2', 'Pending'),
    ('87', 'Avoid monotone - vary pace and emphasis naturally', 'P2', 'Pending'),
    ('88', 'Create script structure templates (hook, body, close)', 'P2', 'Pending'),
    ('89', 'Define attention span guidelines', 'P3', 'Pending'),
    ('90', 'Build engagement hooks library', 'P3', 'Pending'),
    ('91', 'Implement boredom prevention', 'P3', 'Pending'),
]

for item in morning_signal:
    row = table2.add_row().cells
    for i, val in enumerate(item):
        row[i].text = val

doc.add_paragraph()

# Section 3: Content Quality SME Panel
doc.add_heading('Section 3: Content Quality SME Panel (12 items)', level=1)

table3 = doc.add_table(rows=1, cols=4)
table3.style = 'Table Grid'
hdr3 = table3.rows[0].cells
for i, h in enumerate(['#', 'Action', 'Priority', 'Status']):
    hdr3[i].text = h
    hdr3[i].paragraphs[0].runs[0].bold = True

content_quality = [
    ('92', 'Script Writer SME - review all scripts before generation', 'P2', 'Pending'),
    ('93', 'Film Director SME - advise on visual pacing and imagery', 'P2', 'Pending'),
    ('94', 'Voice Director SME - guide tone and delivery', 'P2', 'Pending'),
    ('95', 'Engagement Expert SME - optimize for attention retention', 'P2', 'Pending'),
    ('96', 'Optimize ElevenLabs voice settings for CEPHO tone', 'P2', 'Pending'),
    ('97', 'Implement dynamic pacing in voice scripts', 'P3', 'Pending'),
    ('98', 'Add emphasis markers for key points', 'P3', 'Pending'),
    ('99', 'Create voice style presets (briefing, urgent, casual)', 'P3', 'Pending'),
    ('100', 'Test and refine voice output quality', 'P2', 'Pending'),
    ('101', 'Define visual pacing guidelines', 'P3', 'Pending'),
    ('102', 'Create imagery prompts for key moments', 'P3', 'Pending'),
    ('103', 'Build transition and visual hook library', 'P3', 'Pending'),
]

for item in content_quality:
    row = table3.add_row().cells
    for i, val in enumerate(item):
        row[i].text = val

doc.add_paragraph()

# Section 4: Quality Assurance System
doc.add_heading('Section 4: Quality Assurance System (10 items)', level=1)

table4 = doc.add_table(rows=1, cols=4)
table4.style = 'Table Grid'
hdr4 = table4.rows[0].cells
for i, h in enumerate(['#', 'Action', 'Priority', 'Status']):
    hdr4[i].text = h
    hdr4[i].paragraphs[0].runs[0].bold = True

qa_system = [
    ('104', 'Define engagement metrics for each format', 'P2', 'Pending'),
    ('105', 'Build feedback loop for continuous improvement', 'P2', 'Pending'),
    ('106', 'Track attention retention data', 'P3', 'Pending'),
    ('107', 'A/B test different approaches', 'P3', 'Pending'),
    ('108', 'COS reviews all outputs against quality standards', 'P1', 'Pending'),
    ('109', 'Daily Signal (voice, video, PDF) quality check', 'P1', 'Pending'),
    ('110', 'Project briefings quality check', 'P2', 'Pending'),
    ('111', 'SME consultations quality check', 'P2', 'Pending'),
    ('112', 'Reports and presentations quality check', 'P2', 'Pending'),
    ('113', 'All external communications quality check', 'P1', 'Pending'),
]

for item in qa_system:
    row = table4.add_row().cells
    for i, val in enumerate(item):
        row[i].text = val

doc.add_paragraph()

# Section 5: Digital Twin Questionnaire System
doc.add_heading('Section 5: Digital Twin Questionnaire System (10 items)', level=1)

table5 = doc.add_table(rows=1, cols=4)
table5.style = 'Table Grid'
hdr5 = table5.rows[0].cells
for i, h in enumerate(['#', 'Action', 'Priority', 'Status']):
    hdr5[i].text = h
    hdr5[i].paragraphs[0].runs[0].bold = True

dt_questionnaire = [
    ('114', 'Add 200-question Digital Twin questionnaire to onboarding', 'P1', 'Pending'),
    ('115', 'Create database table to store questionnaire responses', 'P1', 'Pending'),
    ('116', 'Build questionnaire UI component with progress tracking', 'P1', 'Pending'),
    ('117', 'Split questionnaire into digestible sections', 'P2', 'Pending'),
    ('118', 'Store responses and calculate user profile scores', 'P2', 'Pending'),
    ('119', 'Generate Digital Twin profile from questionnaire responses', 'P1', 'Pending'),
    ('120', 'Show COS Understanding Level based on completion', 'P2', 'Pending'),
    ('121', 'Allow users to update answers over time', 'P3', 'Pending'),
    ('122', 'Use questionnaire data to personalize COS behavior', 'P1', 'Pending'),
    ('123', 'Make questionnaire completion part of Getting to 100 score', 'P2', 'Pending'),
]

for item in dt_questionnaire:
    row = table5.add_row().cells
    for i, val in enumerate(item):
        row[i].text = val

doc.add_paragraph()

# Section 6: Chief of Staff KPI Dashboard
doc.add_heading('Section 6: Chief of Staff KPI Dashboard (6 items)', level=1)

table6 = doc.add_table(rows=1, cols=4)
table6.style = 'Table Grid'
hdr6 = table6.rows[0].cells
for i, h in enumerate(['#', 'Action', 'Priority', 'Status']):
    hdr6[i].text = h
    hdr6[i].paragraphs[0].runs[0].bold = True

cos_kpi = [
    ('124', 'Create ChiefOfStaffKPIReport component', 'P1', 'Pending'),
    ('125', 'Build SME Performance Heatmap with color-coded scoring', 'P1', 'Pending'),
    ('126', 'Create Business Area Scoring Table with all SME rankings', 'P1', 'Pending'),
    ('127', 'Implement Enhancement Recommendations engine', 'P1', 'Pending'),
    ('128', 'Add export functionality for report (PDF/CSV)', 'P2', 'Pending'),
    ('129', 'Integrate report into Chief of Staff page', 'P1', 'Pending'),
]

for item in cos_kpi:
    row = table6.add_row().cells
    for i, val in enumerate(item):
        row[i].text = val

doc.add_paragraph()

# Section 7: Process Documentation
doc.add_heading('Section 7: Process Documentation (10 items)', level=1)

table7 = doc.add_table(rows=1, cols=4)
table7.style = 'Table Grid'
hdr7 = table7.rows[0].cells
for i, h in enumerate(['#', 'Action', 'Priority', 'Status']):
    hdr7[i].text = h
    hdr7[i].paragraphs[0].runs[0].bold = True

process_docs = [
    ('130', 'Build process flow document showing each step visually', 'P2', 'Pending'),
    ('131', 'Add decision points and branches in the flow', 'P2', 'Pending'),
    ('132', 'Add imagery cues for video scripts', 'P3', 'Pending'),
    ('133', 'Implement attention retention techniques', 'P3', 'Pending'),
    ('134', 'Add B-roll and visual variety guidance', 'P3', 'Pending'),
    ('135', 'Daily Signal application', 'P2', 'Pending'),
    ('136', 'Project briefings application', 'P2', 'Pending'),
    ('137', 'SME consultations application', 'P2', 'Pending'),
    ('138', 'Reports and presentations application', 'P2', 'Pending'),
    ('139', 'All external communications application', 'P2', 'Pending'),
]

for item in process_docs:
    row = table7.add_row().cells
    for i, val in enumerate(item):
        row[i].text = val

doc.add_paragraph()

# Section 8: Additional Items
doc.add_heading('Section 8: Additional Items (7 items)', level=1)

table8 = doc.add_table(rows=1, cols=4)
table8.style = 'Table Grid'
hdr8 = table8.rows[0].cells
for i, h in enumerate(['#', 'Action', 'Priority', 'Status']):
    hdr8[i].text = h
    hdr8[i].paragraphs[0].runs[0].bold = True

additional = [
    ('143', 'Active learning mode with screen observation', 'P3', 'Pending'),
    ('144', 'Label and thread management for Gmail', 'P3', 'Pending'),
    ('145', 'Push notifications for new emails', 'P2', 'Pending'),
    ('146', 'Real-time webhook for new emails', 'P2', 'Pending'),
    ('147', 'Custom domain setup', 'P2', 'Pending'),
]

for item in additional:
    row = table8.add_row().cells
    for i, val in enumerate(item):
        row[i].text = val

doc.add_paragraph()

# Summary Table
doc.add_heading('Summary by Section', level=1)

summary_table = doc.add_table(rows=1, cols=4)
summary_table.style = 'Table Grid'
sum_hdr = summary_table.rows[0].cells
for i, h in enumerate(['Section', 'Items', 'P1', 'P2/P3']):
    sum_hdr[i].text = h
    sum_hdr[i].paragraphs[0].runs[0].bold = True

sections = [
    ('SME Panel Actions', '20', '-', '20'),
    ('Morning Signal System', '32', '11', '21'),
    ('Content Quality SME Panel', '12', '0', '12'),
    ('Quality Assurance System', '10', '3', '7'),
    ('Digital Twin Questionnaire', '10', '5', '5'),
    ('Chief of Staff KPI Dashboard', '6', '5', '1'),
    ('Process Documentation', '10', '0', '10'),
    ('Additional Items', '7', '0', '7'),
    ('TOTAL', '107', '24', '83'),
]

for section in sections:
    row = summary_table.add_row().cells
    for i, val in enumerate(section):
        row[i].text = val
    if section[0] == 'TOTAL':
        for cell in row:
            cell.paragraphs[0].runs[0].bold = True

doc.add_paragraph()

# Footer
footer = doc.add_paragraph()
footer.add_run('Report Generated: ').bold = True
footer.add_run('18 January 2026 | 13:30 GMT+4\n')
footer.add_run('Logged by: ').bold = True
footer.add_run('Chief of Staff')

doc.save('/home/ubuntu/the-brain/SME_ACTIONS_TODAY_18JAN2026.docx')
print('Word document created successfully with 107 items')
